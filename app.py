
# import pdfplumber
# import docx
# from fuzzywuzzy import fuzz
# import json
# import requests
# import streamlit as st 
# import main
# # Replace with your actual Groq API key
# GROQ_API_KEY = "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"
# GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"


# # Function Definitions (same as before)
# def extract_text_from_pdf(pdf_path):
#     text = ""
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() + "\n" if page.extract_text() else ""
#     return text


# def extract_text_from_docx(docx_path):
#     doc = docx.Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])


# def fetch_job_description(role, company):
#     prompt = f"""
#     Generate a detailed job description for the role of {role} at {company}.
#     Include key responsibilities, required skills, and qualifications.
#     Return the job description as a plain text string.
#     """
#     response = query_groq_api(prompt)
#     return response.strip() if response else "Job description not available."


# def query_groq_api(prompt):
#     headers = {
#         "Authorization": f"Bearer {GROQ_API_KEY}",
#         "Content-Type": "application/json"
#     }
#     data = {
#         "model": "qwen-2.5-32b",  # Replace with a valid Groq model
#         "messages": [{"role": "user", "content": prompt}],
#         "max_tokens": 500
#     }

#     try:
#         response = requests.post(GROQ_API_URL, headers=headers, json=data)
#         response.raise_for_status()
#         return response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
#     except requests.exceptions.RequestException as e:
#         print(f"Error querying Groq API: {e}")
#         return None


# def extract_skills_from_resume(resume_text):
#     prompt = f"""
#     Please extract strictly the key technical skills from the following resume text.
#     Only return the skills in a list format, like this: ["skill1", "skill2", "skill3", ...].
#     Resume Text: {resume_text}
#     Return only a list of skills in JSON format.
#     """
#     response = query_groq_api(prompt)
#     try:
#         skills = json.loads(response)
#         return skills
#     except json.JSONDecodeError:
#         return []


# def extract_skills_from_jd(jd_text):
#     prompt = f"""
#     Please extract strictly the key technical skills from the following job description text.
#     Only return the skills in a list format, like this: ["skill1", "skill2", "skill3", ...].
#     jd_text: {jd_text}
#     Return only a list of skills in JSON format.
#     """
#     response = query_groq_api(prompt)
#     try:
#         jd_skills = json.loads(response)
#         if isinstance(jd_skills, dict) and 'skills' in jd_skills:
#             return jd_skills['skills']
#         return jd_skills
#     except json.JSONDecodeError:
#         return []


# def compare_skills(resume_skills, jd_skills):
#     if isinstance(resume_skills, list) and isinstance(resume_skills[0], str):
#         resume_skill_list = resume_skills
#     elif isinstance(resume_skills, list) and isinstance(resume_skills[0], dict):
#         resume_skill_list = []
#         for item in resume_skills:
#             resume_skill_list.extend(item.get('skills', []))
#     else:
#         raise ValueError("Invalid format for resume_skills.")

#     if isinstance(jd_skills, list) and isinstance(jd_skills[0], str):
#         jd_skill_list = jd_skills
#     elif isinstance(jd_skills, list) and isinstance(jd_skills[0], dict):
#         jd_skill_list = []
#         for item in jd_skills:
#             jd_skill_list.extend(item.get('skills', []))
#     else:
#         raise ValueError("Invalid format for jd_skills.")
    
#     resume_skill_set = set(resume_skill_list)
#     jd_skill_set = set(jd_skill_list)
    
#     similarity = fuzz.partial_ratio(" ".join(resume_skill_set), " ".join(jd_skill_set))
#     missing_skills = list(jd_skill_set - resume_skill_set)
    
#     return similarity, missing_skills


# def generate_interview_level(similarity_score):
#     if similarity_score > 80:
#         level = "Advanced"
#     elif similarity_score > 50:
#         level = "Intermediate"
#     else:
#         level = "Beginner"
#     return level


# def analyze_resume_and_jd(resume_text, jd_text):
#     resume_skills = extract_skills_from_resume(resume_text)
#     jd_skills = extract_skills_from_jd(jd_text)
#     similarity_score, missing_skills = compare_skills(resume_skills, jd_skills)
#     interview_level = generate_interview_level(similarity_score)
#     return {
#         "similarity_score": similarity_score,
#         "missing_skills": missing_skills,
#         "interview_level": interview_level
#     }


# # Streamlit App Layout and Styling
# def main():
#     st.set_page_config(page_title="Mock Interview Analysis", layout="wide")

#     # Add custom CSS for styling
#     st.markdown("""
#         <style>
#         /* Main background */
#         .stApp {
#             background: linear-gradient(to right, #cbe7ff, #a6ddff, #89c9ff);
#             font-family: 'Arial', sans-serif;
#             color: white;
#             padding: 20px;
#         }
#         /* Title */
#         h1 {
#             color: white;
#             font-size: 36px;
#             text-align: center;
#             margin-bottom: 20px;
#         }
#         .stTextInput input, .stFileUploader label {
#             background-color: rgba(255, 255, 255, 0.95);
#             border-radius: 12px;
#             padding: 12px;
#             border: 1px solid rgba(0, 170, 255, 0.2);
#             color: #2C3E50;
#             box-shadow: 0 4px 8px rgba(0, 0, 255, 0.1);
#         }
#         .stTextInput input:focus {
#             border: 2px solid #3fa9f5;
#             outline: none;
#         }

#         /* Input fields */
#         # .stTextInput input, .stFileUploader label {
#         #     background-color: rgba(255, 255, 255, 0.9);
#         #     border-radius: 10px;
#         #     padding: 10px;
#         #     color: #2C3E50;
#         # }
#         /* Button */
#         # .stButton button {
#         #     font-size: 18px;
#         #     padding: 12px 25px;
#         #     background-color: #ff6f61;
#         #     color: white;
#         #     border-radius: 10px;
#         #     border: none;
#         #     cursor: pointer;
#         #     width: 100%;
#         #     transition: background-color 0.3s ease;
#         # }
#         # .stButton button:hover {
#         #     background-color: #ff3b2f;
#         # }
#         .stButton button {
#             font-size: 18px;
#             padding: 12px 25px;
#             background: linear-gradient(to right, #68c2ff, #3fa9f5);
#             color: white;
#             border-radius: 12px;
#             border: none;
#             cursor: pointer;
#             width: 100%;
#             box-shadow: 2px 2px 10px rgba(0, 170, 255, 0.3);
#             transition: all 0.3s ease;
#         }
#         .stButton button:hover {
#             background: linear-gradient(to right, #3fa9f5, #68c2ff);
#             transform: scale(1.05);
#         }


#                 /* Error message */
#         .stAlert {
#             color: #ff3b2f;
#             font-weight: bold;
#             margin-top: 10px;
#         }
#         /* Result box */
#         .result-box {
#             background-color: rgba(255, 255, 255, 0.9);
#             padding: 20px;
#             border-radius: 10px;
#             margin-top: 20px;
#             box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
#             color: #2C3E50;
#         }
      
#         </style>
#     """, unsafe_allow_html=True)

    
#     st.title("Mock Interview Preparation")
#     st.write("---")

#     # Centered container for inputs
#     with st.container():
#         st.markdown("<div class='centered'>", unsafe_allow_html=True)

#         # File Upload and Job Role Input
#         uploaded_file = st.file_uploader("Upload your resume (PDF/DOCX)", type=["pdf", "docx"])
#         role = st.text_input("Job Role", key="role")
#         company = st.text_input("Company Name", key="company")

#         # Button for generating interview
#         if st.button("Generate Interview", key="generate_interview", help="Click to generate interview questions"):
#             if not uploaded_file:
#                 st.error("Please upload your resume.")
#             elif not role:
#                 st.error("Please enter the job role.")
#             elif not company:
#                 st.error("Please enter the company name.")
#             else:
#                 with st.spinner("Analyzing your resume and generating interview questions..."):
#                     if uploaded_file.name.endswith(".pdf"):
#                         resume_text = extract_text_from_pdf(uploaded_file)
#                     elif uploaded_file.name.endswith(".docx"):
#                         resume_text = extract_text_from_docx(uploaded_file)
#                     else:
#                         st.error("Unsupported file format. Use PDF or DOCX.")
                    
#                     # Fetch job description
#                     job_description = fetch_job_description(role, company)

#                     # Store data in session state
#                     st.session_state.resume_text = resume_text
#                     st.session_state.job_description = job_description

#                     # Navigate to the second page
#                     st.session_state.page = "main"
#                     # st.experimental_rerun()
#                     st.rerun()
#         st.markdown("</div>", unsafe_allow_html=True) 

# if __name__ == "__main__":
#     main() 
# import pdfplumber
# import docx
# from fuzzywuzzy import fuzz
# import json
# import requests
# import streamlit as st

# # Replace with your actual Groq API key
# GROQ_API_KEY = "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"
# GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"

# # Function Definitions (same as before)
# def extract_text_from_pdf(pdf_path):
#     text = ""
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() + "\n" if page.extract_text() else ""
#     return text

# def extract_text_from_docx(docx_path):
#     doc = docx.Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def fetch_job_description(role, company):
#     prompt = f"""
#     Generate a detailed job description for the role of {role} at {company}.
#     Include key responsibilities, required skills, and qualifications.
#     Return the job description as a plain text string.
#     """
#     response = query_groq_api(prompt)
#     return response.strip() if response else "Job description not available."

# def query_groq_api(prompt):
#     headers = {
#         "Authorization": f"Bearer {GROQ_API_KEY}",
#         "Content-Type": "application/json"
#     }
#     data = {
#         "model": "qwen-2.5-32b",  # Replace with a valid Groq model
#         "messages": [{"role": "user", "content": prompt}],
#         "max_tokens": 500
#     }

#     try:
#         response = requests.post(GROQ_API_URL, headers=headers, json=data)
#         response.raise_for_status()
#         return response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
#     except requests.exceptions.RequestException as e:
#         print(f"Error querying Groq API: {e}")
#         return None

# def extract_skills_from_resume(resume_text):
#     prompt = f"""
#     Please extract strictly the key technical skills from the following resume text.
#     Only return the skills in a list format, like this: ["skill1", "skill2", "skill3", ...].
#     Resume Text: {resume_text}
#     Return only a list of skills in JSON format.
#     """
#     response = query_groq_api(prompt)
#     try:
#         skills = json.loads(response)
#         return skills
#     except json.JSONDecodeError:
#         return []

# def extract_skills_from_jd(jd_text):
#     prompt = f"""
#     Please extract strictly the key technical skills from the following job description text.
#     Only return the skills in a list format, like this: ["skill1", "skill2", "skill3", ...].
#     jd_text: {jd_text}
#     Return only a list of skills in JSON format.
#     """
#     response = query_groq_api(prompt)
#     try:
#         jd_skills = json.loads(response)
#         if isinstance(jd_skills, dict) and 'skills' in jd_skills:
#             return jd_skills['skills']
#         return jd_skills
#     except json.JSONDecodeError:
#         return []

# def compare_skills(resume_skills, jd_skills):
#     if isinstance(resume_skills, list) and isinstance(resume_skills[0], str):
#         resume_skill_list = resume_skills
#     elif isinstance(resume_skills, list) and isinstance(resume_skills[0], dict):
#         resume_skill_list = []
#         for item in resume_skills:
#             resume_skill_list.extend(item.get('skills', []))
#     else:
#         raise ValueError("Invalid format for resume_skills.")

#     if isinstance(jd_skills, list) and isinstance(jd_skills[0], str):
#         jd_skill_list = jd_skills
#     elif isinstance(jd_skills, list) and isinstance(jd_skills[0], dict):
#         jd_skill_list = []
#         for item in jd_skills:
#             jd_skill_list.extend(item.get('skills', []))
#     else:
#         raise ValueError("Invalid format for jd_skills.")
    
#     resume_skill_set = set(resume_skill_list)
#     jd_skill_set = set(jd_skill_list)
    
#     similarity = fuzz.partial_ratio(" ".join(resume_skill_set), " ".join(jd_skill_set))
#     missing_skills = list(jd_skill_set - resume_skill_set)
    
#     return similarity, missing_skills

# def generate_interview_level(similarity_score):
#     if similarity_score > 80:
#         level = "Advanced"
#     elif similarity_score > 50:
#         level = "Intermediate"
#     else:
#         level = "Beginner"
#     return level

# def analyze_resume_and_jd(resume_text, jd_text):
#     resume_skills = extract_skills_from_resume(resume_text)
#     jd_skills = extract_skills_from_jd(jd_text)
#     similarity_score, missing_skills = compare_skills(resume_skills, jd_skills)
#     interview_level = generate_interview_level(similarity_score)
#     return {
#         "similarity_score": similarity_score,
#         "missing_skills": missing_skills,
#         "interview_level": interview_level
#     }

# # Streamlit App Layout and Styling
# def main():
#     st.set_page_config(page_title="Mock Interview Analysis", layout="wide")

#     # Add custom CSS for styling
#     st.markdown("""
#         <style>
#         /* Main background */
#         .stApp {
#             background: linear-gradient(to right, #cbe7ff, #a6ddff, #89c9ff);
#             font-family: 'Arial', sans-serif;
#             color: white;
#             padding: 20px;
#         }
#         /* Title */
#         h1 {
#             color: white;
#             font-size: 36px;
#             text-align: center;
#             margin-bottom: 20px;
#         }
#         .stTextInput input, .stFileUploader label {
#             background-color: rgba(255, 255, 255, 0.95);
#             border-radius: 12px;
#             padding: 12px;
#             border: 1px solid rgba(0, 170, 255, 0.2);
#             color: #2C3E50;
#             box-shadow: 0 4px 8px rgba(0, 0, 255, 0.1);
#         }
#         .stTextInput input:focus {
#             border: 2px solid #3fa9f5;
#             outline: none;
#         }

#         /* Button */
#         .stButton button {
#             font-size: 18px;
#             padding: 12px 25px;
#             background: linear-gradient(to right, #68c2ff, #3fa9f5);
#             color: white;
#             border-radius: 12px;
#             border: none;
#             cursor: pointer;
#             width: 100%;
#             box-shadow: 2px 2px 10px rgba(0, 170, 255, 0.3);
#             transition: all 0.3s ease;
#         }
#         .stButton button:hover {
#             background: linear-gradient(to right, #3fa9f5, #68c2ff);
#             transform: scale(1.05);
#         }

#         /* Error message */
#         .stAlert {
#             color: #ff3b2f;
#             font-weight: bold;
#             margin-top: 10px;
#         }
#         /* Result box */
#         .result-box {
#             background-color: rgba(255, 255, 255, 0.9);
#             padding: 20px;
#             border-radius: 10px;
#             margin-top: 20px;
#             box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
#             color: #2C3E50;
#         }
#         </style>
#     """, unsafe_allow_html=True)

#     st.title("Mock Interview Preparation")
#     st.write("---")

#     # Centered container for inputs
#     with st.container():
#         st.markdown("<div class='centered'>", unsafe_allow_html=True)

#         # File Upload and Job Role Input
#         uploaded_file = st.file_uploader("Upload your resume (PDF/DOCX)", type=["pdf", "docx"])
#         role = st.text_input("Job Role", key="role")
#         company = st.text_input("Company Name", key="company")

#         # Button for generating interview
#         if st.button("Generate Interview", key="generate_interview", help="Click to generate interview questions"):
#             if not uploaded_file:
#                 st.error("Please upload your resume.")
#             elif not role:
#                 st.error("Please enter the job role.")
#             elif not company:
#                 st.error("Please enter the company name.")
#             else:
#                 with st.spinner("Analyzing your resume and generating interview questions..."):
#                     if uploaded_file.name.endswith(".pdf"):
#                         resume_text = extract_text_from_pdf(uploaded_file)
#                     elif uploaded_file.name.endswith(".docx"):
#                         resume_text = extract_text_from_docx(uploaded_file)
#                     else:
#                         st.error("Unsupported file format. Use PDF or DOCX.")
                    
#                     # Fetch job description
#                     job_description = fetch_job_description(role, company)

#                     # Store data in session state
#                     st.session_state.resume_text = resume_text
#                     st.session_state.job_description = job_description

#                     # Navigate to the second page
#                     st.session_state.page = "main"
#                     st.rerun()  # Use st.rerun() to refresh the app and navigate to the next page

#         st.markdown("</div>", unsafe_allow_html=True)

# if __name__ == "__main__":
#     main() 
import pdfplumber
import docx
import json
import requests
import streamlit as st
import main
# Replace with your actual Groq API key
GROQ_API_KEY = "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"
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"

# Function Definitions (same as before)
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n" if page.extract_text() else ""
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def fetch_job_description(role, company):
    prompt = f"""
    Generate a detailed job description for the role of {role} at {company}.
    Include key responsibilities, required skills, and qualifications.
    Return the job description as a plain text string.
    """
    response = query_groq_api(prompt)
    return response.strip() if response else "Job description not available."

def query_groq_api(prompt):
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "qwen-2.5-32b",  # Replace with a valid Groq model
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 500
    }

    try:
        response = requests.post(GROQ_API_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
    except requests.exceptions.RequestException as e:
        print(f"Error querying Groq API: {e}")
        return None

# Streamlit App Layout and Styling
def mainn():
    st.set_page_config(page_title="Mock Interview Analysis", layout="wide")

    # Add custom CSS for styling
    st.markdown("""
        <style>
        /* Main background */
        .stApp {
            background: linear-gradient(to right, #cbe7ff, #a6ddff, #89c9ff);
            font-family: 'Arial', sans-serif;
            color: white;
            padding: 20px;
        }
        /* Title */
        h1 {
            color: white;
            font-size: 36px;
            text-align: center;
            margin-bottom: 20px;
        }
        .stTextInput input, .stFileUploader label {
            background-color: rgba(255, 255, 255, 0.95);
            border-radius: 12px;
            padding: 12px;
            border: 1px solid rgba(0, 170, 255, 0.2);
            color: #2C3E50;
            box-shadow: 0 4px 8px rgba(0, 0, 255, 0.1);
        }
        .stTextInput input:focus {
            border: 2px solid #3fa9f5;
            outline: none;
        }

        /* Button */
        .stButton button {
            font-size: 18px;
            padding: 12px 25px;
            background: linear-gradient(to right, #68c2ff, #3fa9f5);
            color: white;
            border-radius: 12px;
            border: none;
            cursor: pointer;
            width: 100%;
            box-shadow: 2px 2px 10px rgba(0, 170, 255, 0.3);
            transition: all 0.3s ease;
        }
        .stButton button:hover {
            background: linear-gradient(to right, #3fa9f5, #68c2ff);
            transform: scale(1.05);
        }

        /* Error message */
        .stAlert {
            color: #ff3b2f;
            font-weight: bold;
            margin-top: 10px;
        }
        /* Result box */
        .result-box {
            background-color: rgba(255, 255, 255, 0.9);
            padding: 20px;
            border-radius: 10px;
            margin-top: 20px;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
            color: #2C3E50;
        }
        </style>
    """, unsafe_allow_html=True)

    # Initialize session state for page navigation
    if "page" not in st.session_state:
        st.session_state.page = "home"

    # Home Page
    if st.session_state.page == "home":
        st.title("Mock Interview Preparation")
        st.write("---")

        # Centered container for inputs
        with st.container():
            st.markdown("<div class='centered'>", unsafe_allow_html=True)

            # File Upload and Job Role Input
            uploaded_file = st.file_uploader("Upload your resume (PDF/DOCX)", type=["pdf", "docx"])
            role = st.text_input("Job Role", key="role")
            company = st.text_input("Company Name", key="company")

            # Button for generating interview
            if st.button("Generate Interview", key="generate_interview", help="Click to generate interview questions"):
                if not uploaded_file:
                    st.error("Please upload your resume.")
                elif not role:
                    st.error("Please enter the job role.")
                elif not company:
                    st.error("Please enter the company name.")
                else:
                    with st.spinner("Analyzing your resume and generating interview questions..."):
                        if uploaded_file.name.endswith(".pdf"):
                            resume_text = extract_text_from_pdf(uploaded_file)
                        elif uploaded_file.name.endswith(".docx"):
                            resume_text = extract_text_from_docx(uploaded_file)
                        else:
                            st.error("Unsupported file format. Use PDF or DOCX.")
                        
                        # Fetch job description
                        job_description = fetch_job_description(role, company)

                        # Store data in session state
                        st.session_state.resume_text = resume_text
                        st.session_state.job_description = job_description

                        # Navigate to the second page
                        st.session_state.page = "main"
                        st.rerun()  # Refresh the app to load the second page

            st.markdown("</div>", unsafe_allow_html=True)

    # Main Page (Interview Questions)
    elif st.session_state.page == "main":
         # Import the main.py file
        main.main()  # Call the main function from main.py

if __name__ == "__main__":
    mainn()